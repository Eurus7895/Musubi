#!/usr/bin/env python3
"""Convert Markdown source to PDF or DOCX via pandoc or weasyprint + python-docx."""

import json
import subprocess
import sys
import tempfile
from pathlib import Path

TIMEOUT = 120


def find_pandoc() -> bool:
    try:
        subprocess.run(["pandoc", "--version"], capture_output=True, timeout=5, shell=False)
        return True
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False


def render_with_pandoc(source: str, title: str, fmt: str, out_path: Path) -> dict:
    with tempfile.NamedTemporaryFile(suffix=".md", mode="w", encoding="utf-8", delete=False) as f:
        f.write(f"---\ntitle: \"{title}\"\n---\n\n{source}")
        src = f.name

    pandoc_fmt = "pdf" if fmt == "pdf" else "docx"
    result = subprocess.run(
        ["pandoc", src, "-o", str(out_path), "--standalone"],
        capture_output=True,
        text=True,
        timeout=TIMEOUT,
        shell=False,
    )
    Path(src).unlink(missing_ok=True)

    if result.returncode != 0:
        return {"ok": False, "error": result.stderr.strip() or result.stdout.strip()}
    return {"ok": True}


def render_docx_python(source: str, title: str, out_path: Path) -> dict:
    try:
        from docx import Document  # type: ignore[import]
        from docx.shared import Pt  # type: ignore[import]
    except ImportError:
        return {"ok": False, "error": "python-docx not installed — run: pip install python-docx"}

    doc = Document()
    doc.add_heading(title, level=0)

    for line in source.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped:
            doc.add_paragraph(stripped)

    doc.save(str(out_path))
    return {"ok": True}


def main() -> None:
    try:
        payload = json.loads(sys.stdin.read())
        source = payload.get("source", "").strip()
        title = payload.get("title", "Document")
        fmt = payload.get("format", "pdf").lower()
    except (json.JSONDecodeError, KeyError) as exc:
        print(json.dumps({"ok": False, "error": str(exc)}))
        sys.exit(1)

    if not source:
        print(json.dumps({"ok": False, "error": "source is required"}))
        sys.exit(1)

    if fmt not in ("pdf", "docx"):
        print(json.dumps({"ok": False, "error": f"unsupported format: {fmt!r} — use 'pdf' or 'docx'"}))
        sys.exit(1)

    suffix = ".pdf" if fmt == "pdf" else ".docx"
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as f:
        out_path = Path(f.name)

    # prefer pandoc; fall back to python-docx for docx
    if find_pandoc():
        result = render_with_pandoc(source, title, fmt, out_path)
    elif fmt == "docx":
        result = render_docx_python(source, title, out_path)
    else:
        result = {"ok": False, "error": "pandoc not found — install pandoc or use format 'docx' with python-docx"}

    if not result["ok"]:
        out_path.unlink(missing_ok=True)
        print(json.dumps(result))
        sys.exit(1)

    size = out_path.stat().st_size
    print(json.dumps({"ok": True, "path": str(out_path), "format": fmt, "size_bytes": size}))


if __name__ == "__main__":
    main()
